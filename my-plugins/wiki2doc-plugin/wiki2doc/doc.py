""" class to create a word document """

import re
import os
import tempfile
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
#from trac.mimeview import Context
from trac.mimeview.api import RenderingContext
from trac.util.text import to_unicode
from trac.wiki.formatter import HtmlFormatter
from .helpers import\
add_hyperlink,\
check_table_row_length,\
check_string,\
create_list,\
filter_wiki_text,\
find_hyperlinks,\
get_base_url,\
get_link_name,\
get_wiki_specname,\
insert_image,\
merge_table,\
process_blockquote,\
remove_forward_slash,\
table_font_size

class Doc(object): # pylint: disable=too-many-public-methods
    """ class to create a document in MS Word """

    def __init__(self, args):
        self.document = Document(args[0])
        self.env = args[1]
        self.wiki2doc = args[2]
        self.req = args[3]
        
        self.report_type = None

    def save(self, path):
        """ save docx to path """
        self.document.save(path)

    def get_content(self):
         """ save docx to path """
         _, out = tempfile.mkstemp()
         self.save(out)
         with open(out) as filehndl:
             content = filehndl.read()
         os.unlink(out)
         return content
  
    def get_paragraph_after_regex(self, regex):
         """ helper function to be used before insert_paragraph_before() """
  
         regex = re.compile(regex)
         found = False
         idx = 0
         while not found and idx < len(self.document.paragraphs):
             par = self.document.paragraphs[idx]
             match = regex.match(par.text)
             if match:
                 found = True
             else:
                 idx += 1
         idx += 1
         if found:
             if idx < len(self.document.paragraphs):
                 return self.document.paragraphs[idx]
         return self.document.add_paragraph()
  
    def get_merge_row(self, params):
         """ for a given table data, analyses
             the data and finds merged cells.
             params = [idr, row, table_row_length,
                       col_size, row_length, table,
                       task_id, spec_name]"""
  
         merge_row = []
         col = 0
         pos = 0
         start = 0
         end = 0
  
         for idx, item in enumerate(params[1]):
             for idy, value in enumerate(item):
                 if check_table_row_length(params[3],
                                           params[4]):
                     value = filter_wiki_text(value)
                     #args = [table, paragraph,
                     #        text, task_id, spec_name]
                     args = [params[5],
                             params[5].rows[params[0]].cells[col].paragraphs[0],
                             value,
                             params[6],
                             params[7]]
                     #args[2].rows[args[0]].cells[args[1]].paragraphs[0]
                     params[5], _ = self.filter_hyperlinks(args)
                 else:
                     params[2].add(False)
                 col += 1
                 start = pos
                 end = pos + len(item)
                 if idy == 0 and value == '' and idx < len(params[1])-1:
                     merge_row.append([start, end])
                 elif idy == 0 and value == '' and \
                     idx == len(params[1])-1 and len(item) > 1:
                     merge_row.append([start, end-1])
                 elif idy == 0 and col == params[4] - 1 and \
                     len(merge_row) == 0:
                     merge_row.append([])
             pos += len(item)
  
         return (params[5], params[2], merge_row)
  
  
    def find_merged_cells(self, args):
        """ for a given table data, analyses
            the data and finds merged cells.
            args = [data,
                    table,
                    col_size,
                    task_id,
                    spec_name]"""
        merge_list = []
        table_row_length = set()
        for idr, row in enumerate(args[0]):
            row_length = 0
  
            for item in row:
                row_length += len(item)
            params = [idr,
                      row,
                      table_row_length,
                      args[2],
                      row_length,
                      args[1],
                      args[3],
                      args[4]]
            args[1], table_row_length, merge_row = self.get_merge_row(params)
  
            merge_list.append(merge_row)
            merge_row = []
            row_length = 0
  
            if idr < len(args[0])-1:
                args[1].add_row()
  
        if len(list(table_row_length)) > 0:
            page_path = get_base_url(self.req) +\
                'Coconut/event/wiki/' + args[4]
            self.autorep.errorlog.append((
                "There might be an extra pipe || in the wikitext of" +\
                " a table that needs to be removed. Number of" +\
                " columns in each row must match including merged" +\
                " cells! Check the following table with a:" +\
                " header: {}".format(args[0][0]),
                get_base_url(self.req) + \
                'Coconut/task/ticket/' + str(args[3]),
                page_path))
        return (args[1], merge_list)
  
    def append_table(self, data, task_id, spec_name):
        """ for a given table data, this function analyzes the table,
            looks for cells to be merged inside the text as described
            below. Creates the table with the values first, then calls
            merge_table method to merge the cells if they need to be merged,
            then returns the table.
  
            Wiki Markup:
            || 1 || 2 || 3 ||
            |||| 1-2 || 3 ||
            || 1 |||| 2-3 ||
            |||||| 1-2-3 ||
  
            Display:
            ---- --- ----
            | 1 | 2 | 3 |
            ---- --- ----
            | 1-2   | 3 |
            -------- ----
            | 1 | 2-3   |
            -------------
            | 1-2-3     |
            -------------"""
        merge_list = []
        col_size = 0
  
        for item in data[0]:
            col_size += len(item)
  
        table = self.document.add_table(rows=1, cols=col_size)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER # pylint: disable=no-member
        args = [data,
                table,
                col_size,
                task_id,
                spec_name]
        table, merge_list = self.find_merged_cells(args)
  
        if len(data) == len(merge_list):
            table = merge_table(table, merge_list)
            table = table_font_size(table, 8)
            return table
        else:
            self.autorep.errorlog.append(
                "Merge cell list length and table length does not match." +\
                "Please check the merged cells in: {} \n".format(data[0]),
                0,
                'None')
  
    def insert_table(self, paragraph, data, task_id, spec_name):
        """ insert table """
  
        table = self.append_table(data, task_id, spec_name)
        table.style = 'Table Grid'
  
        if paragraph is not None:
            new = paragraph.insert_paragraph_before()
            # _p and _tbl are protected and therefore not documented,
            # but otherwise it is impossible to insert a table,
            # and it would only be possible to append it to the
            # end of the document
            new._p.addnext(table._tbl) # pylint: disable=protected-access
  
    def errorlog_missing_page(self, task_id, missing_spec, in_spec):
        """ Errorlog a wiki page that
            does not exist. """
        missing_spec = remove_forward_slash(missing_spec)
        self.autorep.errorlog.append(\
            ("Specified link does not exist. Please check the " +\
             "full path and the name of the following link:'{}']".format(\
                                                           missing_spec),
             get_base_url(self.req) + \
                 'Coconut/task/ticket/' + str(task_id),
             get_base_url(self.req) + \
             'Coconut/event/wiki/' + in_spec))
  
    def get_wiki_hyperlink(self, spec_name, hyper):
        """ returns the wiki page hyperlink path for
            another page that is under same parent
            path. See regex_id 4. in find_hyperlinks.
  
            Example wikipage: http://10.45.43.145:8000/Coconut/
            event/wiki/APO/IP006/Dummy-APO-Database/IP006-APO-Spec-Sill
  
            Example reference from inside the link above.
            [[Dummy-APO-Database/GPD/Material_Strength| MS-GPD]]
            [[Dummy-APO-Database/GPD/Metallic_Joint| MBJ-GPD]]
            [[/GPD/Material_Strength| MS-GPD]]
            [[GPD/Material_Strength| MS-GPD]]
            [[/IP006/Dummy-APO-Database/GPD/Material_Strength| MS-GPD]]
            [[IP006/Dummy-APO-Database/GPD/Material_Strength| MS-GPD]]
            [[IP006/Dummy-APO-Database/GPD/Material_Strength]]
  
            This works because both "IP006-APO-Spec-Sill" and
            "GPD/Material_Strength" are under:
  
            http://10.45.43.145:8000/Coconut/
            event/wiki/APO/IP006/Dummy-APO-Database/
            """
  
        given_path = remove_forward_slash(hyper[1]) + hyper[2]
        given_path_list = given_path.split("/")
        full_wiki_path = get_base_url(self.req) +\
            "Coconut/event/wiki/" + spec_name
        full_path_list = full_wiki_path.split("/")
        protocol = full_path_list[0]
        full_path_list = full_path_list[2:]
        list_index = []
        hyperlink = ''
  
        for i, item in enumerate(full_path_list):
            if item in set(given_path_list):
                list_index.append(i)
  
        if len(list_index) > 0:
            full_path_list = full_path_list[:list_index[0]]
        elif len(list_index) == 0:
            full_path_list = full_path_list[:-1]
  
        mod_full_path = ''
  
        for item in full_path_list:
            mod_full_path += item + "/"
  
        hyperlink = protocol + "//" + mod_full_path + given_path
  
        return hyperlink
  
    def get_link_path(self, task_id, spec_name, regex_id, hyper):
        """ for a given hypermatch this function
            returns the hyperlink path."""
  
        hyperlink = ''
  
        if regex_id == 4:
            another_child_spec_name = get_wiki_specname(spec_name, hyper)
            page = \
                self.autorep.get_wikipage(
                    remove_forward_slash(another_child_spec_name))
            if not page:
                self.errorlog_missing_page(task_id,
                                           another_child_spec_name,
                                           spec_name)
            hyperlink = self.get_wiki_hyperlink(spec_name, hyper)
        elif hyper[1] == '/wiki/':
            page = self.autorep.get_wikipage(hyper[2])
            if not page:
                self.errorlog_missing_page(task_id, hyper[2], spec_name)
            hyperlink = get_base_url(self.req) +\
                'Coconut/event/' + remove_forward_slash(hyper[1]) +\
                hyper[2]
        elif hyper[1] == 'e:/wiki/':
            page = self.autorep.get_wikipage(hyper[2])
            if not page:
                self.errorlog_missing_page(task_id, hyper[2], spec_name)
            hyperlink = get_base_url(self.req) +\
            'Coconut/event/wiki/' + hyper[2]
        elif hyper[1] == 'wiki:':
            page = \
                self.autorep.get_wikipage(remove_forward_slash(hyper[2]))
            if not page:
                self.errorlog_missing_page(task_id, hyper[2], spec_name)
            hyperlink = get_base_url(self.req) +\
                'Coconut/event/wiki/' + hyper[2]
        elif hyper[1] == 'r:#':
            hyperlink = get_base_url(self.req) +\
                'Coconut/req/ticket/' + hyper[2]
        else:
            hyperlink = hyper[1] + hyper[2]
  
        return hyperlink
  
    def get_hyperlink(self, task_id, spec_name, regex_id, hyper):
        """ for a given hypermatch this function
            returns the hyperlink and the link name."""
  
        link_name = get_link_name(hyper)
        hyperlink = self.get_link_path(task_id, spec_name, regex_id, hyper)
  
        return (hyperlink, link_name)
  
    def filter_hyperlinks(self, args):
        """ for a given paragraph text or a table text,
            this function filters the table text
            and returns the table data
            args = [table,
                    paragraph,
                    text,
                    task_id,
                    spec_name]"""
  
        context = Context.from_request(self.req, 'wiki')
        regex_id, hypermatches = find_hyperlinks(args[2])
        hyperlink = ''
        if len(hypermatches) > 0:
            link_name = ''
            rest = hypermatches.pop()
            for hyper in hypermatches:
                flt_text = ''
                if self.report_type == 'ADC':
                    #args[2].rows[args[0]].cells[args[1]].\
                    #    paragraphs[0].add_run(hyper[0])
                    wiki = process_blockquote(check_string(hyper[0]))
                    self.parse_html(args, context, wiki)
                    hyperlink, link_name = self.get_hyperlink(args[3],
                                                              args[4],
                                                              regex_id,
                                                              hyper)
                    if hyperlink == None:
                        break
                    add_hyperlink(args[1],
                                  hyperlink,
                                  link_name,
                                  '0000FF',
                                  True)
                elif self.report_type == 'SAR':
                    flt_text = flt_text + hyper[0]
                    #args[2].rows[args[0]].cells[args[1]].\
                    #    paragraphs[0].add_run(flt_text)
                    wiki = process_blockquote(check_string(flt_text))
                    self.parse_html(args, context, wiki)
            #args[2].rows[args[0]].cells[args[1]].paragraphs[0].add_run(rest)
            wiki = process_blockquote(check_string(rest))
            self.parse_html(args, context, wiki)
        else:
            wiki = process_blockquote(check_string(args[2]))
            self.parse_html(args, context, wiki)
            #args[2].rows[args[0]].cells[args[1]].text = \
            #        unicode(args[3], "utf-8")
        return (args[0], hypermatches)
  
    def parse_html(self, args, context, wiki):
        """ Parse html string to docx
        args[1] = paragraph,
        context,
        wiki,
        args[3] = task_id,
        args[4] = spec_name"""
  
        try:
            html_code = HtmlFormatter(self.envs['event'],
                                      context,
                                      wiki).generate()
            DocumentHTMLParser(self.document, args[1], html_code)
            return html_code
        except AttributeError:
            self.autorep.errorlog.append(
                ("HtmlFormatter could not parse" +\
                 " the following wikitext: {}".format(wiki),
                 get_base_url(self.req) + \
                 'Coconut/task/ticket/' + str(args[3]),
                 get_base_url(self.req) + \
                 'Coconut/event/wiki/' + args[4]))
  
    def find_sections(self, params):
        """ Given paragraph location and sections data,
            inserts section text, if found images and
            if found tables.
            params = [i,
                      paragraph,
                      sections,
                      text,
                      spec_images]"""
        img_filename = None
        img_path = None
        wiki_filter = \
            [re.compile(r'\s*\[\[Image\((.*(\.jpg|\.png|\.gif))\)\]\]\s*'),
             re.compile(r'\s*\[\[Table\((.*)\.tbl\)\]\]\s*'),
             re.compile(r'\s*(=+)\s*(\d+\.){1,}\d*(.*)'),
             re.compile(r'\s*\[\s*=\#Table(\d+)\s*\]\s*'),
             re.compile(r'\s*\[\s*=\#Fig(\d+)\s*\]\s*'),
             re.compile(r'\s*\*\s*(.*)')]
#         image = re.compile(r'\s*\[\[Image\((.*(\.jpg|\.png|\.gif))\)\]\]\s*')
#         anchor = re.compile(r'\s*\[\[Table\((.*)\.tbl\)\]\]\s*')
#         section = re.compile(r'\s*(=+)\s*(\d+\.){1,}\d*(.*)')
#         tbl = re.compile(r'\s*\[\s*=\#Table(\d+)\s*\]\s*')
#         fig = re.compile(r'\s*\[\s*=\#Fig(\d+)\s*\]\s*')
  
        for line in params[3].splitlines():
            line = to_unicode(line)
#             img_match = wiki_filter[0].match(line)
#             anc_match = wiki_filter[1].match(line)
#             sec_match = wiki_filter[2].match(line)
#             tbl_match = wiki_filter[3].match(line)
#             fig_match = wiki_filter[4].match(line)
            if wiki_filter[0].match(line):
                img_filename = to_unicode(wiki_filter[0].match(line).group(1))
                for key, value in params[4].iteritems():
                    if key == img_filename:
                        img_path = value
                        # if you want to include the image name
                        # insert the code below
                        # params[1].insert_paragraph_before(line)
                        insert_image(params[1], img_path)
            elif wiki_filter[1].match(line):
                self.get_table(params,
                               to_unicode(wiki_filter[1].match(line).group(1)))
            elif wiki_filter[2].match(line):
                style_key = 'Heading' +\
                            ' ' + \
                            str(len(wiki_filter[2].match(line).group(1))+1)
                params[1].insert_paragraph_before(\
                    to_unicode(wiki_filter[2].match(line).group(3).strip()),
                    style=style_key)
            elif wiki_filter[3].match(line):
                line = 'Table' + ' ' + str(wiki_filter[3].match(line).group(1))
                line = to_unicode(line)
                params[1].insert_paragraph_before(line, style='Caption')
            elif wiki_filter[4].match(line):
                line = 'Figure' + ' ' + str(wiki_filter[4].match(line).group(1))
                line = to_unicode(line)
                params[1].insert_paragraph_before(line, style='Caption')
            elif wiki_filter[5].match(line):
                line = str(wiki_filter[5].match(line).group(1))
                line = to_unicode(line)
                paragraph = create_list(\
                    params[1].insert_paragraph_before(text=' ',
                                                      style='List Bullet'))
                line = filter_wiki_text(line)
                args = [None,
                        paragraph,
                        line,
                        params[2][params[0]][0],
                        params[2][params[0]][1]]
                self.filter_hyperlinks(args)
            else:
                line = filter_wiki_text(line)
                args = [None,
                        params[1].insert_paragraph_before(),
                        line,
                        params[2][params[0]][0],
                        params[2][params[0]][1]]
                self.filter_hyperlinks(args)
  
    def get_table(self, params, match_group):
        """ Gets table information from
            list of sections and calls
            insert_table method.
            example_sections = [
            [2,
             'Specname1',
             'consetetur sadipscing elitr, sed diam nonumy eirmod tempor\n' +\
             '[[Image(Image1.jpg)]]\n' +\
             'invidunt ut labore et dolore magna \n' +\
             'aliquyam erat, sed diam \n' +\
             '[[Image(Image2.jpg)]]\nvoluptua.\n',
             {'Image1.jpg': u'/tmp/trac-tempenv-WbQieJ/files/attachments/' +\
              'wiki/bdc/bdc726f49cd502d4306404b090a5ddd13bb7dc0e/98c78c01' +\
              'ccdb21a78fd4f561e980ccd4d3a5a685.jpg',
              'Image2.jpg': u'/tmp/trac-tempenv-WbQieJ/files/attachments/' +\
              'wiki/bdc/bdc726f49cd502d4306404b090a5ddd13bb7dc0e/e8385af6' +\
              'dfec928ba93ae7b6ccdc2c5f2fcb89f8.jpg'},
             {'Table_11': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}],
               ...]]
            params = [i,
                      paragraph,
                      sections,
                      text,
                      spec_images]"""
  
        if params[2][params[0]][4]:
            for value in params[2][params[0]][4]:
                if value == match_group:
                    table = params[2][params[0]][4][value]
                    self.insert_table(params[1],
                                      table,
                                      params[2][params[0]][0],
                                      params[2][params[0]][1])
  
    def insert_analysed_apos_table(self, paragraph, sections):
        """ Given paragraph location and sections data,
            creates a table that contains analysed
            apo information."""
  
        col_names = ('ITEM', 'ANALYSE APO TASK No', 'APO TASK NAME', 'REMARK')
  
        table = self.document.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER # pylint: disable=no-member
        table.style = 'Table Grid'
  
        for idx, name in enumerate(col_names):
            prg = table.rows[0].cells[idx].paragraphs[0]
            run = prg.add_run(name)
            run.font.name = 'Arial Black'
            run.bold = True
            run.italic = True
  
        for idx in range(len(sections)):
            table.add_row()
            prg = table.rows[idx+1].cells[0].paragraphs[0]
            run = prg.add_run(str(idx+1))
            run.bold = True
            for idy, cell in enumerate(table.rows[idx+1].cells):
                if idy == 1:
                    cell.paragraphs[0].add_run(str(sections[idx][0]))
                elif idy == 2:
                    cell.paragraphs[0].add_run(str(sections[idx][1]))
  
        table = table_font_size(table, 8)
  
        if paragraph is not None:
            new = paragraph.insert_paragraph_before()
            # _p and _tbl are protected and therefore not documented,
            # but otherwise it is impossible to insert a table,
            # and it would only be possible to append it to the
            # end of the document
            new._p.addnext(table._tbl) # pylint: disable=protected-access
  
    def insert_section(self, paragraph, sections, level):
        """ Given paragraph location and sections data,
            inserts section text, if found images and
            if found tables."""
  
        spec_images = {}
  
        for i in range(len(sections)):
            text = sections[i][2]
            spec_images.update(sections[i][3])
            apo_spec = str(sections[i][0]) + ", " + sections[i][1]
            if self.report_type == 'SAR':
                apo_spec_path = sections[i][1].split('/')
                apo_spec = apo_spec_path[-1]
            style_key = 'Heading '+ str(level)
            paragraph.insert_paragraph_before(apo_spec, style=style_key)
            if text is not None:
                params = [i, paragraph, sections, text, spec_images]
                self.find_sections(params)

